#!/usr/bin/env python3
"""
Gera docs/relatorio-modelo-custos.docx a partir do conteúdo do hotsite.

Produz um relatório de negócio em estilo corporativo neutro (Calibri,
preto/cinza), com título e as sete seções do index.html. O conteúdo é
mantido neste script; ao alterar o hotsite, atualize os textos aqui e
reexecute: ``python scripts/build-report-docx.py``.

Dependência: python-docx (pip install python-docx).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "relatorio-modelo-custos.docx"

BASE_FONT = "Calibri"
COLOR_HEADING = RGBColor(0x33, 0x33, 0x33)
COLOR_SUBTITLE = RGBColor(0x55, 0x55, 0x55)
COLOR_NOTE_SHADE = "F2F2F2"


def configure_styles(document: Document) -> None:
    """Define fonte base e cores neutras para o corpo e títulos."""
    normal = document.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    for level, size in (("Title", 22), ("Heading 1", 16), ("Heading 2", 14)):
        style = document.styles[level]
        style.font.name = BASE_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = COLOR_HEADING
        style.font.bold = True


def shade_paragraph(paragraph, fill: str = COLOR_NOTE_SHADE) -> None:
    """Aplica sombreamento de fundo a um parágrafo (usado nas notas)."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_intro(document: Document, text: str) -> None:
    """Adiciona parágrafo de corpo justificado."""
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_labeled_item(document: Document, label: str, text: str) -> None:
    """Adiciona item de lista com rótulo em negrito seguido da descrição."""
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(text)


def add_numbered_item(document: Document, label: str, text: str) -> None:
    """Adiciona item numerado com rótulo em negrito seguido da descrição."""
    paragraph = document.add_paragraph(style="List Number")
    run = paragraph.add_run(f"{label} — ")
    run.bold = True
    paragraph.add_run(text)


def add_note(document: Document, label: str, text: str) -> None:
    """Adiciona nota destacada (sombreada) com rótulo em negrito."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.right_indent = Pt(12)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(text)
    shade_paragraph(paragraph)


def add_timeline_table(document: Document, rows: list[tuple[str, str]]) -> None:
    """Adiciona a linha do tempo como tabela de duas colunas."""
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    header[0].paragraphs[0].add_run("Período").bold = True
    header[1].paragraphs[0].add_run("Iniciativa").bold = True
    for period, initiative in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(period).bold = True
        cells[1].paragraphs[0].add_run(initiative)


def build() -> None:
    """Monta o relatório completo e grava o arquivo .docx."""
    document = Document()
    configure_styles(document)

    # 1. Título + abertura (#hero)
    document.add_heading("Modelo de Custos da Transformação Digital", level=0)

    subtitle = document.add_paragraph()
    sub_run = subtitle.add_run(
        "Mensuração automatizada, contínua e orientada por dados para a "
        "Administração Pública."
    )
    sub_run.bold = True
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = COLOR_SUBTITLE

    add_intro(
        document,
        "Uma metodologia desenvolvida pela Secretaria de Governo Digital (SGD) "
        "para identificar e mensurar os impactos econômicos, sociais e "
        "ambientais da digitalização de serviços públicos.",
    )

    # 2. Breve Histórico (#historico)
    document.add_heading("Breve Histórico", level=1)
    add_intro(
        document,
        "Modelos de custos consistem em estruturas analíticas e matemáticas "
        "destinadas à identificação e mensuração dos custos associados a "
        "normas, atividades e processos, sendo amplamente utilizados em "
        "diversos países. Entre os principais referenciais internacionais, "
        "destaca-se o Standard Cost Model (SCM), criado na Holanda no início "
        "dos anos 2000 para medir o custo da burocracia imposta pela regulação. "
        "A experiência foi bem-sucedida e o modelo passou a ser adotado pela "
        "Comissão Europeia e recomendado pela OCDE como padrão de mensuração.",
    )
    add_intro(
        document,
        "No Brasil, com o objetivo de mensurar a economia de custos decorrente "
        "da digitalização de serviços públicos, a Secretaria de Governo Digital "
        "(SGD) desenvolveu, a partir de 2019, diversas iniciativas:",
    )
    add_timeline_table(
        document,
        [
            ("2019", "Início das iniciativas"),
            ("2020", "Calculadora de Interoperabilidade (CONECTAGOV)"),
            ("2022", "Guia Prático – Modelo de Custos de Serviços Públicos (SGD)"),
            (
                "2025",
                "Relatório Final da Proposta de Modelo de Custos de Serviços "
                "Públicos (SGD/UFMG)",
            ),
            (
                "Atual",
                "Referencial unificado — Modelo de Custos Aplicado à "
                "Digitalização",
            ),
        ],
    )
    add_intro(
        document,
        "Como cada iniciativa abordou a temática sob perspectivas distintas, "
        "somada à necessidade de um estudo mais abrangente e atualizado, o "
        "Modelo atual consolida os elementos relevantes das iniciativas "
        "anteriores em uma metodologia única, integrando com os dados "
        "estruturados do Portal Gov.br e possibilitando a replicabilidade por "
        "Estados e Municípios.",
    )

    # 3. Automação e Evolução Orgânica (#integracao)
    document.add_heading("Automação e Evolução Orgânica", level=1)
    add_intro(
        document,
        "O principal diferencial deste Modelo é a utilização do Portal Gov.br "
        "como fonte primária e direta de dados. Diferente de abordagens "
        "anteriores que dependiam de estimativas manuais enviadas por gestores, "
        "este sistema extrai e processa dados estruturados na origem.",
    )
    add_labeled_item(
        document,
        "Sem formulários manuais",
        "Extingue a necessidade de coleta de dados semiestruturados ou baseados "
        "na memória gerencial de órgãos públicos.",
    )
    add_labeled_item(
        document,
        "Precisão evolutiva",
        "A acurácia das estimativas cresce proporcionalmente ao amadurecimento "
        "e à melhoria das descrições dos serviços no Portal Gov.br.",
    )
    add_labeled_item(
        document,
        "Atualização periódica",
        "O modelo monitora de forma contínua o volume de solicitações, gerando "
        "um histórico dinâmico de economia para o Estado e para a sociedade.",
    )
    add_note(
        document,
        "Fontes",
        "os dados estruturados vêm da página de serviços do Governo Federal "
        "(https://www.gov.br/pt-br/servicos) e de dados complementares do "
        "próprio Portal Gov.br.",
    )

    # 4. O Serviço como Pilar Central (#servico)
    document.add_heading("O Serviço como Pilar Central", level=1)
    add_intro(
        document,
        "O Modelo estabelece o Serviço Público como unidade básica de "
        "importância e cálculo. Por ser a unidade central de interação no "
        "Gov.br, focar no serviço garante escalabilidade e permite que cada "
        "entrega à população seja mensurada de forma individualizada. Para "
        "isso, o Modelo adota um conjunto padronizado de etapas aplicáveis a "
        "qualquer serviço, composto por:",
    )
    add_numbered_item(
        document,
        "Buscar Informações",
        "Atividade do usuário para compreender os requisitos do serviço. No "
        "cenário digital, elimina o atendimento presencial e o deslocamento "
        "físico, sendo contabilizada apenas uma vez por serviço.",
    )
    add_numbered_item(
        document,
        "Coletar Documentos",
        "Reunião de insumos necessários pelo usuário. O modelo adota a premissa "
        "prudente de que o cidadão já possui os documentos, medindo "
        "exclusivamente o tempo de localização interna.",
    )
    add_numbered_item(
        document,
        "Solicitar Serviço",
        "Entrega de dados e preenchimento de formulários. O tempo varia "
        "conforme a quantidade de documentos exigidos e a presença de "
        "validações automatizadas via integração de dados.",
    )
    add_numbered_item(
        document,
        "Receber Serviço",
        "Obtenção do resultado final de forma online, extinguindo a necessidade "
        "de interações físicas e entregas em suporte de papel.",
    )
    add_note(
        document,
        "Nota",
        "cada serviço poderá conter uma ou mais ocorrências dessas etapas ou "
        "mesmo não contemplar determinada etapa, conforme sua natureza. Essa "
        "padronização garante flexibilidade e uniformidade metodológica, "
        "permitindo que os tempos de execução variem conforme as "
        "especificidades do serviço. As etapas são classificadas "
        "automaticamente nas macroetapas Solicitar e Receber.",
    )

    # 5. Engenharia e Operacionalização do Modelo (#motor)
    document.add_heading("Engenharia e Operacionalização do Modelo", level=1)
    add_intro(
        document,
        "O processamento percorre cinco estágios funcionais, da captura dos "
        "dados brutos à consolidação visual dos indicadores.",
    )
    add_numbered_item(
        document,
        "Extração automatizada",
        "Captura de dados brutos tabulares e estruturados em formato JSON a "
        "partir do Portal Gov.br.",
    )
    add_numbered_item(
        document,
        "Tratamento léxico",
        "Limpeza de caracteres e normalização de textos inseridos nos serviços "
        "do Portal Gov.br.",
    )
    add_numbered_item(
        document,
        "Classificação preditiva",
        "Análise linguística feita por algoritmos (IA) que processam os verbos "
        "das descrições dos serviços do Portal Gov.br, categorizando as ações "
        "entre as macroetapas Solicitar e Receber.",
    )
    add_numbered_item(
        document,
        "Mensuração multidimensional",
        "Aplicação de fórmulas integrando tempo, quantidade de documentos e "
        "volume de uso detectado via API de Avaliação da Qualidade dos Serviços "
        "do Portal Gov.br, além de variáveis globais e atualizadas que "
        "complementam as estimativas.",
    )
    add_numbered_item(
        document,
        "Painel estratégico",
        "Consolidação de todas as métricas em painéis de BI, garantindo a "
        "transparência ativa e o acesso aos dados.",
    )

    # 6. Mensuração Além dos Custos Financeiros (#indicadores)
    document.add_heading("Mensuração Além dos Custos Financeiros", level=1)
    add_intro(
        document,
        "O modelo cruza dados específicos de serviços com variáveis globais do "
        "ecossistema. Todo o cálculo baseia-se no princípio da prudência "
        "(estimativas conservadoras), garantindo que os resultados econômicos "
        "apresentados sejam sempre subestimados para evitar efeito de "
        "superavaliação.",
    )
    add_labeled_item(
        document,
        "Perspectiva econômica",
        "Ganho de eficiência administrativa do Governo e redução de custos de "
        "deslocamento, transporte e oportunidade para a Sociedade.",
    )
    add_labeled_item(
        document,
        "Impacto social",
        "Mensuração do ganho de bem-estar social em horas livres devolvidas ao "
        "cidadão, redução da congestão urbana e diminuição da assimetria "
        "informacional.",
    )
    add_labeled_item(
        document,
        "Impacto ambiental",
        "Indicadores físicos que medem a mitigação do desmatamento (redução do "
        "uso de papel), diminuição de emissões de CO₂ na atmosfera e controle "
        "do consumo energético institucional.",
    )

    # 7. Tomada de Decisão Baseada em Evidências (#transparencia)
    document.add_heading("Tomada de Decisão Baseada em Evidências", level=1)
    add_intro(
        document,
        "Os resultados do modelo são disponibilizados em um painel interativo "
        "(Power BI), permitindo segmentações dinâmicas. Este referencial "
        "técnico foi desenhado para ser totalmente replicável por entes "
        "subnacionais (estados e municípios), estimulando a homogeneidade de "
        "critérios, fortalecendo a Governança Digital em toda a Federação e "
        "promovendo a tomada de decisão da Política Pública baseada em "
        "evidências.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(f"Gerado: {OUT.relative_to(ROOT)} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
